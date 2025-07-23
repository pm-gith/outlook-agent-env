import json
import re
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime, timedelta

def classify_email(subject, body):
    text = subject + " " + body
    if re.search(r"\burgent|asap|action required|due|Prachi\b", text, re.I):
        return "Actionable"
    elif re.search(r"automated|noreply|newsletter", text, re.I):
        return "Auto-Generated"
    else:
        return "Informational"
    
def categorize_date(received_time):
    """
    Convert received_time string to 'Today', 'Yesterday', 'This Week', or 'Older'.
    """
    try:
        email_date = datetime.fromisoformat(received_time.replace('Z', '+00:00'))
    except Exception:
        return "Older"
    
    now = datetime.now(email_date.tzinfo)
    today = now.date()
    email_day = email_date.date()

    if email_day == today:
        return "Received Today"
    elif email_day == today - timedelta(days=1):
        return "Received Yesterday"
    elif email_day >= today - timedelta(days=7):
        return "Received This Week"
    else:
        return "Older than a Week"

def clean_html(html_content):
    """Remove HTML tags and return clean text."""
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text(separator="\n", strip =True)

def summarize_body(body):
    text = clean_html(body)
    lines = text.split("\n")
    highlights = [line.strip() for line in lines if len(line.strip()) > 40][:5]
    return " ".join(highlights)

def process_json_file(filename):
    #with open(os.path.join(FOLDER_PATH, filename), "r", encoding="utf-8") as f:
    with open(filename, "r", encoding="utf-8") as f:
        data = json.load(f)

    emails = data.get("body", [])  # Access the 'body' list

    summary = {"Actionable": [], "Informational": [], "Auto-Generated": []}
    for mail in emails:
        subject = mail.get("email subject", "")
        body = mail.get("email body", "")
        category = classify_email(subject, body)
        short = summarize_body(body)
        summary[category].append(f"📌 {subject}\n🔍 {short}\n")

    return summary

def save_to_word(summary, output_path):
    doc = Document()
    doc.add_heading('Email Summary', level=1)

    for section, items in summary.items():
        doc.add_heading(section, level=2)
        for item in items:
            doc.add_paragraph(item)

    doc.save(output_path)
    print(f"Summary saved to {output_path}")
